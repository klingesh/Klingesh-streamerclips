#!/usr/bin/env python3
"""
Generates a polished, plain-English Word guide for the aptitude topic
"Percentage". It explains the concept in very easy language and covers every
common type of percentage sum with step-by-step worked examples, shortcuts,
and practice questions.

Reuses the document style/approach of generate_branding_guide.py.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Colour palette
# ----------------------------------------------------------------------------
NAVY = RGBColor(0x0C, 0x12, 0x1F)
BLUE = RGBColor(0x1A, 0x68, 0xFF)
CYAN = RGBColor(0x60, 0xD2, 0xFF)
TEAL = RGBColor(0x40, 0xC4, 0xFF)
GREEN = RGBColor(0x1E, 0x8E, 0x3E)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x7A, 0x7A, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    return p


def add_para(text, bold=False, italic=False, color=None, size=11, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_callout(label, text, fill="EAF6FF", label_color=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    r = cell.paragraphs[0].add_run(label)
    r.bold = True
    r.font.color.rgb = label_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(text)
    r2.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_example(title, steps, answer):
    """A worked example box: title, list of step lines, and a bold answer."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "FFF6E5")
    first = True
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    r = p0.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(0xB8, 0x6A, 0x00)
    for line in steps:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.size = Pt(10.5)
    pa = cell.add_paragraph()
    pa.paragraph_format.space_after = Pt(2)
    ra = pa.add_run("Answer: " + answer)
    ra.bold = True
    ra.font.color.rgb = GREEN
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def make_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0C121F")
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        shade = "FFFFFF" if r_i % 2 == 0 else "EAF6FF"
        for i, val in enumerate(row):
            set_cell_bg(cells[i], shade)
            para = cells[i].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
    return table


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9D4E2")
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break():
    doc.add_page_break()


# ============================================================================
# COVER
# ============================================================================
for _ in range(4):
    doc.add_paragraph()

add_para("APTITUDE MADE EASY", bold=True, color=TEAL, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Percentage")
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A Simple, Complete Guide with All Types of Sums")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = BLUE

add_para("Learn percentages from zero \u2014 in plain, everyday language.",
         color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para(
    "This guide explains what a percentage really means, the few formulas you "
    "need, and every common type of percentage problem you will see in exams and "
    "interviews \u2014 each with easy, step-by-step worked examples and practice "
    "questions with answers.",
    italic=True, color=LIGHT, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading("What's Inside", level=1)
toc = [
    "1.  What is a Percentage? (The simple idea)",
    "2.  Converting between Percentages, Fractions & Decimals",
    "3.  Must-Memorise Fraction \u2194 Percentage Table",
    "4.  Type 1: Finding a Percentage of a Number",
    "5.  Type 2: What Percentage is One Number of Another?",
    "6.  Type 3: Percentage Increase and Decrease",
    "7.  Type 4: Finding the Original / Whole Value",
    "8.  Type 5: Successive (Back-to-Back) Percentage Changes",
    "9.  Type 6: Percentage Error & 'More / Less Than' Problems",
    "10. Type 7: Exam Marks (Pass / Fail) Problems",
    "11. Type 8: Income, Expenditure & Savings",
    "12. Type 9: Population, Price & Depreciation (Growth over time)",
    "13. Type 10: Product / Length-Breadth Change Problems",
    "14. Handy Shortcuts & Tricks",
    "15. Common Mistakes to Avoid",
    "16. Practice Set (with Answers)",
]
for t in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(t)

page_break()

# ============================================================================
# 1. WHAT IS A PERCENTAGE
# ============================================================================
add_heading("1. What is a Percentage? (The simple idea)", level=1)
add_para(
    "The word 'percent' simply means 'out of 100'. The symbol for percent is %. "
    "So when we say 25%, we just mean '25 out of 100'.")
add_para(
    "Imagine a chocolate bar cut into 100 equal pieces. If you eat 25 pieces, you "
    "have eaten 25% of the bar. That's the whole idea \u2014 a percentage is just a "
    "way of describing a part of something, where the whole is always treated as 100.")

add_callout(
    "Remember this one line:",
    "x% means x / 100. For example, 40% = 40/100 = 0.40. Once you see a percentage as "
    "'a number divided by 100', everything else becomes easy.")

add_para("A percentage is just another way of writing a fraction or a decimal:", space_after=4)
add_bullet("50% is the same as the fraction 1/2 and the decimal 0.5.")
add_bullet("25% is the same as 1/4 and 0.25.")
add_bullet("100% means the whole thing (all of it). 200% means twice as much.")

add_callout(
    "Why percentages are useful:",
    "They let us compare things fairly. Scoring 45 out of 50 in one test and 80 out of "
    "100 in another is hard to compare directly \u2014 but 90% vs 80% makes the winner "
    "obvious instantly.")

page_break()

# ============================================================================
# 2. CONVERSIONS
# ============================================================================
add_heading("2. Converting between Percentages, Fractions & Decimals", level=1)
add_para("These three are just different outfits for the same number. Switching "
         "between them is the most important basic skill.")

add_heading("a) Percentage \u2192 Fraction", level=2, color=BLUE)
add_para("Write the number over 100, then simplify.")
add_example(
    "Convert 60% to a fraction",
    ["Step 1: 60% = 60/100",
     "Step 2: Simplify by dividing top and bottom by 20 \u2192 3/5"],
    "60% = 3/5")

add_heading("b) Percentage \u2192 Decimal", level=2, color=BLUE)
add_para("Divide by 100 (just move the decimal point two places to the left).")
add_example(
    "Convert 72% to a decimal",
    ["72% = 72 \u00f7 100 = 0.72"],
    "72% = 0.72")

add_heading("c) Fraction \u2192 Percentage", level=2, color=BLUE)
add_para("Multiply the fraction by 100.")
add_example(
    "Convert 3/4 to a percentage",
    ["Step 1: (3/4) \u00d7 100 = 300/4",
     "Step 2: 300 \u00f7 4 = 75"],
    "3/4 = 75%")

add_heading("d) Decimal \u2192 Percentage", level=2, color=BLUE)
add_para("Multiply by 100 (move the decimal point two places to the right).")
add_example(
    "Convert 0.35 to a percentage",
    ["0.35 \u00d7 100 = 35"],
    "0.35 = 35%")

page_break()

# ============================================================================
# 3. MUST-MEMORISE TABLE
# ============================================================================
add_heading("3. Must-Memorise Fraction \u2194 Percentage Table", level=1)
add_para(
    "Learning these by heart is the single biggest time-saver in percentage "
    "questions. When you instantly know that 1/8 = 12.5%, many sums can be solved "
    "in your head in seconds.")

make_table(
    ["Fraction", "Percentage", "Fraction", "Percentage"],
    [
        ["1/1", "100%", "1/8", "12.5%"],
        ["1/2", "50%", "3/8", "37.5%"],
        ["1/3", "33.33%", "5/8", "62.5%"],
        ["2/3", "66.67%", "7/8", "87.5%"],
        ["1/4", "25%", "1/9", "11.11%"],
        ["3/4", "75%", "1/10", "10%"],
        ["1/5", "20%", "1/11", "9.09%"],
        ["2/5", "40%", "1/12", "8.33%"],
        ["1/6", "16.67%", "1/15", "6.67%"],
        ["1/7", "14.28%", "1/20", "5%"],
    ],
)

add_para("", space_after=4)
add_callout(
    "Pro tip:",
    "Notice the pattern: 1/8 = 12.5%, so 3/8 = 3 \u00d7 12.5% = 37.5%. Once you know the "
    "'unit' fraction (like 1/8), you can build the rest by simple multiplication.")

page_break()

# ============================================================================
# 4. TYPE 1
# ============================================================================
add_heading("4. Type 1: Finding a Percentage of a Number", level=1)
add_para("This is the most common percentage sum. The question looks like: "
         "'What is 20% of 150?'")

add_callout(
    "Formula:",
    "x% of N = (x / 100) \u00d7 N.  In short, turn the percent into a fraction or "
    "decimal, then multiply by the number.")

add_example(
    "Example 1: Find 20% of 150",
    ["Step 1: 20% = 20/100 = 0.20",
     "Step 2: 0.20 \u00d7 150 = 30"],
    "30")

add_example(
    "Example 2: A shirt costs Rs. 800. You get a 15% discount. How much is the discount?",
    ["Discount = 15% of 800 = (15/100) \u00d7 800",
     "= 0.15 \u00d7 800 = 120"],
    "Rs. 120 (so you pay 800 \u2212 120 = Rs. 680)")

add_callout(
    "Mental-maths trick:",
    "10% of any number = just move the decimal one place left (10% of 250 = 25). "
    "Then build others: 5% is half of 10%, 20% is double 10%, 1% is 10% of 10%.")

page_break()

# ============================================================================
# 5. TYPE 2
# ============================================================================
add_heading("5. Type 2: What Percentage is One Number of Another?", level=1)
add_para("Here the question is reversed. Example: '30 is what percent of 120?' or "
         "'A student scored 45 out of 60 \u2014 what percent is that?'")

add_callout(
    "Formula:",
    "Percentage = (Part / Whole) \u00d7 100.  The 'whole' is the number that comes "
    "after the word 'of'.")

add_example(
    "Example 1: 30 is what percent of 120?",
    ["Part = 30, Whole = 120",
     "= (30 / 120) \u00d7 100",
     "= 0.25 \u00d7 100 = 25"],
    "30 is 25% of 120")

add_example(
    "Example 2: A student got 45 marks out of 60. What is the percentage?",
    ["= (45 / 60) \u00d7 100",
     "= 0.75 \u00d7 100 = 75"],
    "75%")

add_callout(
    "Watch out:",
    "The order matters! '30 is what % of 120' (= 25%) is very different from "
    "'120 is what % of 30' (= 400%). Always put the 'whole' on the bottom.")

page_break()

# ============================================================================
# 6. TYPE 3
# ============================================================================
add_heading("6. Type 3: Percentage Increase and Decrease", level=1)
add_para("These problems ask by what percent a value went up or down. Examples: a "
         "price rises, a population grows, or salary increases.")

add_callout(
    "Formulas:",
    "% Increase = (Increase / Original) \u00d7 100.    "
    "% Decrease = (Decrease / Original) \u00d7 100.    "
    "Always divide the CHANGE by the ORIGINAL (starting) value.")

add_example(
    "Example 1: A price went from Rs. 200 to Rs. 250. What is the % increase?",
    ["Increase = 250 \u2212 200 = 50",
     "% Increase = (50 / 200) \u00d7 100 = 25"],
    "25% increase")

add_example(
    "Example 2: A salary fell from Rs. 40,000 to Rs. 34,000. Find the % decrease.",
    ["Decrease = 40000 \u2212 34000 = 6000",
     "% Decrease = (6000 / 40000) \u00d7 100 = 15"],
    "15% decrease")

add_heading("Increasing / decreasing a number by a given percent", level=2, color=BLUE)
add_para("To increase N by x%:  N \u00d7 (1 + x/100).  To decrease N by x%:  N \u00d7 (1 \u2212 x/100).")
add_example(
    "Increase 500 by 20%",
    ["New value = 500 \u00d7 (1 + 20/100)",
     "= 500 \u00d7 1.20 = 600"],
    "600")

page_break()

# ============================================================================
# 7. TYPE 4
# ============================================================================
add_heading("7. Type 4: Finding the Original / Whole Value", level=1)
add_para("Sometimes you are given a part and the percentage, and asked to find the "
         "original (the whole). This is 'working backwards'.")

add_callout(
    "Formula:",
    "Whole = (Part \u00d7 100) / Percentage.  Or think: if 20% = 60, then 1% = 3, so "
    "100% = 300.")

add_example(
    "Example 1: 20% of a number is 60. What is the number?",
    ["1% of the number = 60 / 20 = 3",
     "100% = 3 \u00d7 100 = 300"],
    "The number is 300")

add_example(
    "Example 2: After a 10% discount, a phone costs Rs. 13,500. What was the original price?",
    ["After 10% discount, you pay 90% of the price.",
     "So 90% = 13,500",
     "1% = 13,500 / 90 = 150",
     "100% = 150 \u00d7 100 = 15,000"],
    "Original price = Rs. 15,000")

add_callout(
    "Key idea:",
    "When something is reduced by 10%, the new value is 90% of the original \u2014 NOT "
    "100%. Always figure out what percent the given amount represents first.")

page_break()

# ============================================================================
# 8. TYPE 5
# ============================================================================
add_heading("8. Type 5: Successive (Back-to-Back) Percentage Changes", level=1)
add_para("Here a value changes twice (or more) one after another \u2014 for example, a "
         "price increases by 10%, then by 20%. You CANNOT just add 10% + 20% = 30%. "
         "You must apply them one at a time.")

add_callout(
    "Quick formula for two changes (a% then b%):",
    "Net % change = a + b + (a \u00d7 b)/100.  Use + for an increase and \u2212 for a "
    "decrease in the value of a or b.")

add_example(
    "Example 1: A price increases by 10%, then by 20%. Net change?",
    ["Net % = 10 + 20 + (10 \u00d7 20)/100",
     "= 30 + 200/100 = 30 + 2 = 32"],
    "32% increase (not 30%)")

add_example(
    "Example 2: A price rises 20%, then falls 20%. Net change?",
    ["Here a = +20, b = \u221220",
     "Net % = 20 + (\u221220) + (20 \u00d7 \u221220)/100",
     "= 0 \u2212 400/100 = \u22124"],
    "4% net DECREASE (you do not end up where you started!)")

add_callout(
    "Remember:",
    "An x% increase followed by an x% decrease always gives a net LOSS of (x\u00b2/100)%. "
    "That's why up-20% then down-20% loses 4%.")

page_break()

# ============================================================================
# 9. TYPE 6
# ============================================================================
add_heading("9. Type 6: Percentage Error & 'More / Less Than' Problems", level=1)

add_heading("a) 'How much percent more / less' than another", level=2, color=BLUE)
add_para("Always compare against the number that comes after 'than'. That number "
         "goes on the bottom.")
add_example(
    "A is 600, B is 500. A is what % more than B?",
    ["Difference = 600 \u2212 500 = 100",
     "Compare to B (the 'than' number): (100 / 500) \u00d7 100 = 20"],
    "A is 20% more than B")
add_example(
    "Using the same numbers: B is what % less than A?",
    ["Difference = 100, but now compare to A = 600",
     "(100 / 600) \u00d7 100 = 16.67"],
    "B is 16.67% less than A (different answer \u2014 the base changed!)")

add_heading("b) Percentage error", level=2, color=BLUE)
add_para("Used in measurement problems. It tells how far a measured/estimated value "
         "is from the true value.")
add_callout(
    "Formula:",
    "% Error = (|Measured \u2212 Actual| / Actual) \u00d7 100.")
add_example(
    "True length is 50 cm, but it was measured as 48 cm. Find % error.",
    ["Error = |48 \u2212 50| = 2",
     "% Error = (2 / 50) \u00d7 100 = 4"],
    "4% error")

page_break()

# ============================================================================
# 10. TYPE 7
# ============================================================================
add_heading("10. Type 7: Exam Marks (Pass / Fail) Problems", level=1)
add_para("A classic exam favourite. A student needs a certain pass percentage; you "
         "are told by how many marks they passed or failed, and must find the total "
         "or the pass mark.")

add_example(
    "Example 1: A student needs 40% to pass. He scored 150 marks and failed by 30 marks. Find the maximum marks.",
    ["Pass mark = 150 + 30 = 180",
     "This pass mark is 40% of the total.",
     "So 40% of Total = 180",
     "1% = 180 / 40 = 4.5  \u2192  100% = 450"],
    "Maximum marks = 450")

add_example(
    "Example 2: A student scored 35% and failed by 40 marks. Another scored 60% and got 35 marks more than the pass mark. Find max marks.",
    ["Pass mark from student 1 = 35% of T + 40",
     "Pass mark from student 2 = 60% of T \u2212 35",
     "So 35%T + 40 = 60%T \u2212 35",
     "75 = 25%T  \u2192  25% of T = 75  \u2192  T = 300"],
    "Maximum marks = 300")

add_callout(
    "Strategy:",
    "'Failed by 30 marks' means the pass mark = score + 30. 'Passed with 20 marks to "
    "spare' means pass mark = score \u2212 20. Set the pass mark as a % of total, then solve.")

page_break()

# ============================================================================
# 11. TYPE 8
# ============================================================================
add_heading("11. Type 8: Income, Expenditure & Savings", level=1)
add_para("These word problems involve money: someone spends a percentage and saves "
         "the rest. The key relationship is simple.")
add_callout(
    "Key relationship:",
    "Income = Expenditure + Savings.  If you spend x% of income, you save (100 \u2212 x)% "
    "of income.")

add_example(
    "Example 1: A man spends 75% of his income and saves Rs. 5,000. Find his income.",
    ["He saves (100 \u2212 75)% = 25% of income.",
     "25% of income = 5,000",
     "1% = 5,000 / 25 = 200  \u2192  100% = 20,000"],
    "Income = Rs. 20,000")

add_example(
    "Example 2: A person's income increases by 20% but expenses rise by 10%. If earlier income was 10,000 and expense 6,000, find new savings.",
    ["Old savings = 10,000 \u2212 6,000 = 4,000",
     "New income = 10,000 \u00d7 1.20 = 12,000",
     "New expense = 6,000 \u00d7 1.10 = 6,600",
     "New savings = 12,000 \u2212 6,600 = 5,400"],
    "New savings = Rs. 5,400")

page_break()

# ============================================================================
# 12. TYPE 9
# ============================================================================
add_heading("12. Type 9: Population, Price & Depreciation (Growth over time)", level=1)
add_para("When a quantity grows or shrinks by the same percent every year, we apply "
         "the change repeatedly \u2014 just like compound interest.")

add_callout(
    "Formulas:",
    "Growth: Final = P \u00d7 (1 + r/100)^n.    "
    "Decline/Depreciation: Final = P \u00d7 (1 \u2212 r/100)^n.    "
    "Here r = rate per year, n = number of years.")

add_example(
    "Example 1: A town's population is 10,000 and grows 10% per year. Find it after 2 years.",
    ["Year 1: 10,000 \u00d7 1.10 = 11,000",
     "Year 2: 11,000 \u00d7 1.10 = 12,100"],
    "Population = 12,100")

add_example(
    "Example 2: A machine worth Rs. 50,000 depreciates 20% each year. Value after 2 years?",
    ["Year 1: 50,000 \u00d7 0.80 = 40,000",
     "Year 2: 40,000 \u00d7 0.80 = 32,000"],
    "Value = Rs. 32,000")

add_callout(
    "Finding a past value:",
    "To go backwards (find the value n years ago), DIVIDE instead of multiply: "
    "Past = Present / (1 + r/100)^n.")

page_break()

# ============================================================================
# 13. TYPE 10
# ============================================================================
add_heading("13. Type 10: Product / Length-Breadth Change Problems", level=1)
add_para("These involve a quantity made of two things multiplied together \u2014 like "
         "Area = Length \u00d7 Breadth, or Expenditure = Price \u00d7 Quantity. When both "
         "change by a percent, the same successive-change formula applies.")

add_callout(
    "Formula (same as successive change):",
    "Net % change = a + b + (a \u00d7 b)/100, where a and b are the percent changes in "
    "the two factors.")

add_example(
    "Example 1: A rectangle's length increases 20% and breadth increases 30%. Find the % change in area.",
    ["a = 20, b = 30",
     "Net % = 20 + 30 + (20 \u00d7 30)/100",
     "= 50 + 6 = 56"],
    "Area increases by 56%")

add_example(
    "Example 2: Price of sugar rises 25%. By what % must a family cut consumption to keep the bill the same?",
    ["To keep Price \u00d7 Quantity unchanged after a 25% rise:",
     "Required cut % = (rise / (100 + rise)) \u00d7 100",
     "= (25 / 125) \u00d7 100 = 20"],
    "Reduce consumption by 20%")

page_break()

# ============================================================================
# 14. SHORTCUTS
# ============================================================================
add_heading("14. Handy Shortcuts & Tricks", level=1)
add_bullet(" x% of y is always equal to y% of x. So 16% of 25 = 25% of 16 = 4. Use the easier one!", bold_lead="The flip trick \u2013")
add_bullet(" 10% = move decimal one place left. 1% = move two places left. Build everything else from these.", bold_lead="The 10% anchor \u2013")
add_bullet(" An x% increase followed by an x% decrease = a net loss of (x\u00b2/100)%.", bold_lead="Up-then-down \u2013")
add_bullet(" To keep a product the same when one factor rises by r%, reduce the other by (r / (100 + r)) \u00d7 100 %.", bold_lead="Constant product \u2013")
add_bullet(" Convert ugly percentages to known fractions: 12.5% = 1/8, 16.67% = 1/6, 33.33% = 1/3. Multiplying by a fraction is faster.", bold_lead="Fraction swap \u2013")

add_callout(
    "Speed example:",
    "Find 12.5% of 64. Instead of multiplying by 0.125, use 12.5% = 1/8, so 64 \u00f7 8 = 8. "
    "Done in your head!")

# ============================================================================
# 15. COMMON MISTAKES
# ============================================================================
add_heading("15. Common Mistakes to Avoid", level=1)
add_bullet(" Don't add 10% + 20% to get 30%. Successive changes multiply, giving 32%.", bold_lead="Adding successive % \u2013")
add_bullet(" '% more' and '% less' use different bases, so they give different answers. Read carefully.", bold_lead="Wrong base \u2013")
add_bullet(" A 10% discount means the price becomes 90% of original, not 100%. Work out the right base first.", bold_lead="Forgetting the new base \u2013")
add_bullet(" 'x is what % of y' puts y on the bottom. Swapping them flips the answer completely.", bold_lead="Part/Whole mix-up \u2013")
add_bullet(" Convert percent to a fraction/decimal BEFORE multiplying. 20% of 150 is 0.20 \u00d7 150, not 20 \u00d7 150.", bold_lead="Skipping the \u00f7100 \u2013")

page_break()

# ============================================================================
# 16. PRACTICE SET
# ============================================================================
add_heading("16. Practice Set (with Answers)", level=1)
add_para("Try these yourself first, then check the answers at the bottom. They cover "
         "every type from this guide.", space_after=8)

questions = [
    "1.  What is 35% of 240?",
    "2.  18 is what percent of 72?",
    "3.  A price rises from Rs. 80 to Rs. 100. Find the % increase.",
    "4.  30% of a number is 75. Find the number.",
    "5.  A salary of Rs. 25,000 is increased by 12%. Find the new salary.",
    "6.  A value increases by 25% and then decreases by 20%. Net % change?",
    "7.  A is 250 and B is 200. A is what % more than B?",
    "8.  A student needs 33% to pass, scores 120 and fails by 30 marks. Find total marks.",
    "9.  A man spends 80% of income and saves Rs. 4,000. Find his income.",
    "10. A car worth Rs. 6,00,000 depreciates 10% per year. Value after 2 years?",
    "11. Length of a rectangle rises 10% and breadth rises 10%. Find % change in area.",
    "12. Find 37.5% of 96 (use the fraction trick).",
]
for q in questions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.add_run(q)

divider()
add_para("Answers", bold=True, color=GREEN, size=13)
answers = [
    "1. 84      2. 25%      3. 25%      4. 250",
    "5. Rs. 28,000      6. 0% (no net change)      7. 25% more",
    "8. Pass mark = 150; 33% of T = 150 \u2192 T \u2248 454.5 (so 455 marks)",
    "9. Rs. 20,000      10. Rs. 4,86,000      11. 21% increase      12. 36",
]
for a in answers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(a)
    r.font.size = Pt(10.5)

divider()
add_para(
    "Final word: percentages are everywhere \u2014 discounts, marks, salaries, news "
    "headlines. Master the one core idea ('out of 100'), memorise the fraction table, "
    "and practise each type a few times. Very soon you'll solve most percentage sums "
    "in your head. You've got this!",
    italic=True, color=GREY, size=11)

import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out_path = os.path.join(out_dir, "Percentage_Aptitude_Guide.docx")
doc.save(out_path)
print("Percentage aptitude guide created successfully at:", out_path)
