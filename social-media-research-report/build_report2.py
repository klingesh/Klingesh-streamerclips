# -*- coding: utf-8 -*-
"""Builds the Social Media research report (.docx) from REAL analysis (n=286 consenting)."""
import os, shutil
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "/projects/sandbox"
OUT_DIR = "/projects/sandbox/Klingesh-streamerclips/social-media-research-report"
os.makedirs(OUT_DIR, exist_ok=True)

# extra chart: mean influence by income (the significant ANOVA)
df = pd.read_csv(os.path.join(SRC, "real2_scored.csv"))
inc_order = ["Below ₹20,000","₹20,001–₹40,000","₹40,001–₹60,000","₹60,001–₹80,000","Above ₹80,000"]
inc = df["Monthly Income"].astype(str).str.strip()
plt.figure(figsize=(6,4))
df.assign(i=inc).groupby("i")["Influence_Score"].mean().reindex(inc_order).plot(kind="bar", color="#C44E52")
plt.title("Mean Influence Score by Monthly Income"); plt.ylabel("Mean Score"); plt.xticks(rotation=20, ha="right")
plt.tight_layout(); plt.savefig(os.path.join(SRC,"s_score_by_income.png"), dpi=120); plt.close()

for png in ["s_gender.png","s_age.png","s_platform.png","s_income.png","s_purchased.png",
            "s_score_by_age.png","s_score_by_income.png"]:
    shutil.copy(os.path.join(SRC, png), os.path.join(OUT_DIR, png))

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5; st.paragraph_format.space_after = Pt(6)

def pb(): doc.add_page_break()
def heading(t, level=1):
    h = doc.add_heading(t, level=level)
    for r in h.runs:
        r.font.name = "Times New Roman"; r.font.color.rgb = RGBColor(0x1F,0x3B,0x57)
def para(t, bold=False, italic=False, align="justify", size=12):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align=="center" else WD_ALIGN_PARAGRAPH.JUSTIFY
    return p
def bullet(t): doc.add_paragraph(t, style="List Bullet")
def numbered(t): doc.add_paragraph(t, style="List Number")
def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        run=t.rows[0].cells[i].paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(11)
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            run=c[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(11)
    doc.add_paragraph()
def img(name, width=4.6, cap=None):
    doc.add_picture(os.path.join(OUT_DIR,name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap:
        c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=c.add_run(cap); rr.italic=True; rr.font.size=Pt(10)

# ---------------- TITLE
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("IMPACT OF SOCIAL MEDIA ON\nCONSUMER BUYING BEHAVIOUR"); r.bold=True; r.font.size=Pt(18)
doc.add_paragraph()
for line,sz,b in [("A PROJECT REPORT",13,True),
    ("Submitted in partial fulfilment of the requirements",12,False),
    ("for the award of the degree of",12,False),
    ("MASTER OF BUSINESS ADMINISTRATION",13,True),("",12,False),
    ("Submitted by",12,False),("[Your Name]",13,True),("[Register Number]",12,False),("",12,False),
    ("Under the guidance of",12,False),("[Guide Name]",13,True),("",12,False),
    ("[Department / School Name]",12,True),("[University / College Name]",12,True),("June 2026",12,True)]:
    pp=doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=pp.add_run(line); rr.bold=b; rr.font.size=Pt(sz)
pb()

# ---------------- DECLARATION
heading("DECLARATION")
para('I hereby declare that the project report titled "Impact of Social Media on Consumer Buying '
     'Behaviour" submitted by me is a record of original work carried out under the guidance of '
     '[Guide Name]. The findings reported in this study are based on primary data collected through '
     'a structured questionnaire and have not been submitted earlier for the award of any degree, '
     'diploma, or similar title.')
doc.add_paragraph(); doc.add_paragraph(); para("Place: [City]"); para("Date: [Date]")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p.add_run("[Your Name]\n[Register Number]").bold=True
pb()

# ---------------- ACKNOWLEDGEMENT
heading("ACKNOWLEDGEMENT")
para("I express my sincere gratitude to my guide [Guide Name] for the valuable guidance and support "
     "throughout this study. I thank the Head of the Department and the faculty of [Department Name] "
     "for their encouragement. I am grateful to the 321 respondents who participated in the survey, "
     "and especially to the 286 who provided informed consent and complete responses. Finally, I "
     "thank my family and friends for their constant support.")
pb()

# ---------------- ABSTRACT
heading("ABSTRACT")
para("Social media has become a powerful force shaping how consumers discover, evaluate, and "
     "purchase products. This study examines the impact of social media on consumer buying "
     "behaviour. Primary data were collected from 321 respondents through a structured "
     "questionnaire; after applying an informed-consent filter, 286 valid responses were analysed. "
     "Buying-behaviour influence was measured using a 21-statement five-point Likert scale, and the "
     "data were analysed using percentage analysis, an independent-samples t-test, one-way ANOVA, "
     "and correlation. The overall mean influence score was 3.68 on a 5-point scale, indicating "
     "that social media has an above-average influence on buying behaviour, and 77.3% of "
     "respondents reported having purchased a product after seeing it on social media. The t-test "
     "revealed a statistically significant difference in influence between male and female "
     "respondents, with females being more influenced. One-way ANOVA showed a statistically "
     "significant difference across income groups, with lower-income respondents being more "
     "influenced, whereas differences across age groups were not significant. The study concludes "
     "that social media significantly influences buying behaviour, particularly among female and "
     "lower-income consumers, and offers recommendations for marketers and consumers.")
para("Keywords: social media, consumer buying behaviour, influencer marketing, purchase intention, "
     "advertising, Instagram.", italic=True)
pb()

# ---------------- TOC
heading("TABLE OF CONTENTS")
table(["Chapter","Title"], [
    ("Chapter 1","Introduction"),("Chapter 2","Review of Literature"),
    ("Chapter 3","Research Methodology"),("Chapter 4","Data Analysis and Interpretation"),
    ("Chapter 5","Findings, Suggestions, Conclusion and Summary"),
    ("","Appendix I - Questionnaire"),("","Appendix II - Bibliography / References")])
pb()

# ================= CHAPTER 1
heading("CHAPTER 1: INTRODUCTION")
heading("1.1 Background of the Study",2)
para("Over the last decade, social media platforms such as Instagram, YouTube, Facebook, and "
     "WhatsApp have evolved from simple networking tools into powerful commercial ecosystems. "
     "Consumers no longer rely solely on traditional advertising; instead, they discover products "
     "through influencer recommendations, user-generated content, customer reviews, and targeted "
     "advertisements that appear directly in their social feeds. This shift has fundamentally "
     "altered the consumer decision-making journey, from awareness and interest to evaluation and "
     "final purchase.")
para("In India in particular, the rapid growth of affordable smartphones and internet access has "
     "placed social media at the centre of everyday life. For businesses, this represents an "
     "unprecedented opportunity to reach and influence consumers at scale. Understanding how social "
     "media shapes buying behaviour is therefore essential for marketers, businesses, and consumers "
     "alike.")
heading("1.2 Concept of Social Media Marketing",2)
para("Social media marketing refers to the use of social media platforms to promote products and "
     "services, build brand awareness, and engage with customers. It encompasses paid advertising, "
     "influencer collaborations, content marketing, and community engagement. Unlike traditional "
     "media, social media is interactive and allows two-way communication, enabling consumers to "
     "share opinions, post reviews, and influence one another's purchasing decisions.")
heading("1.3 Consumer Buying Behaviour",2)
para("Consumer buying behaviour refers to the decisions and actions of individuals in searching "
     "for, evaluating, purchasing, and using products and services. It is influenced by "
     "psychological, social, cultural, and personal factors. In the context of social media, "
     "factors such as influencer credibility, peer reviews, promotional offers, and visual content "
     "play an increasingly important role in shaping these decisions.")
heading("1.4 Statement of the Problem",2)
para("While social media is widely used for marketing, the extent and nature of its influence on "
     "actual buying behaviour are not uniform across consumers. Marketers need to understand which "
     "factors drive purchases and whether the influence varies across demographic groups such as "
     "gender, age, and income. This study addresses the problem by empirically measuring the "
     "influence of social media on buying behaviour and testing for demographic differences.")
heading("1.5 Objectives of the Study",2)
for o in [
    "To study the social media usage pattern of respondents.",
    "To examine the influence of social media on consumer buying behaviour.",
    "To analyse the role of influencers, advertisements, reviews, and promotions in purchase decisions.",
    "To test whether the influence of social media differs significantly between male and female consumers.",
    "To test whether the influence of social media differs significantly across income groups.",
    "To offer suggestions to marketers and consumers based on the findings."]:
    numbered(o)
heading("1.6 Scope of the Study",2)
para("The study focuses on individual consumers who use social media. It measures the self-reported "
     "influence of social media on their buying behaviour across dimensions such as advertising, "
     "influencers, reviews, promotions, and user-generated content. The study is based on primary "
     "data collected from 321 respondents, of whom 286 provided informed consent and were included "
     "in the analysis. Data were collected in 2026.")
heading("1.7 Need and Significance of the Study",2)
para("The findings of this study are valuable to marketers seeking to design effective social media "
     "campaigns, to businesses allocating advertising budgets, and to consumers who wish to make "
     "more informed and conscious purchase decisions. By identifying which demographic groups are "
     "more influenced, the study also helps in targeting marketing efforts more precisely.")
heading("1.8 Limitations of the Study",2)
for l in [
    "The study relies on self-reported perceptions, which may differ from actual purchase behaviour.",
    "The sample is dominated by younger respondents and students, which may limit generalisability.",
    "The internal reliability of the scale was moderate (Cronbach's alpha = 0.53), as discussed in Chapter 4.",
    "The study is cross-sectional and does not capture changes in behaviour over time."]:
    bullet(l)
pb()

# ================= CHAPTER 2
heading("CHAPTER 2: REVIEW OF LITERATURE")
para("This chapter reviews studies published in the past ten years (2016-2026) on social media "
     "marketing, influencers, advertising, and their impact on consumer buying behaviour, to "
     "establish the theoretical foundation and identify the research gap.")
lit = [
 ("Mehta and Kulkarni (2020)","found that increased interaction with personalised social media "
  "advertisements influences a consumer's psyche and behaviour to make a purchase even in the "
  "absence of an internal stimulus.","https://www.researchgate.net/publication/352291263"),
 ("A study on social media marketing activities and customer intentions (2021)","explored social "
  "media marketing activities and their impact on continuance, participation, and purchase "
  "intentions, highlighting the mediating roles of social identification and satisfaction.",
  "https://www.frontiersin.org/journals/psychology/articles/10.3389/fpsyg.2021.808525/full"),
 ("A study on social media advertisements and purchase intention (2021)","examined how social "
  "media advertisements influence consumer purchase intention and confirmed a significant positive "
  "relationship.","https://www.tandfonline.com/doi/full/10.1080/23311975.2021.2000697"),
 ("A meta-analysis of social media influencers (2023)","concluded that the entertainment value of "
  "influencers has the strongest association with customer engagement, while influencer credibility "
  "affects purchase intention more than any other attribute.","https://www.mdpi.com/2071-1050/15/3/2744"),
 ("A study on Instagram influencer marketing and online impulse buying (2021)","revealed that the "
  "perceived social media marketing activities of Instagram influencers significantly and "
  "positively affect the source-credibility dimensions of attractiveness, expertise, and "
  "trustworthiness.","https://www.researchgate.net/publication/353828933"),
 ("A study on Instagram influencers as opinion leaders (2021)","found, based on a sample of 223 "
  "followers, that originality, quality, and quantity of content are essential for a user to be "
  "perceived as an opinion leader who shapes purchase behaviour.",
  "https://www.researchgate.net/publication/357111885"),
 ("A study on social commerce through influencers - Indian perspective (2022)","reported that the "
  "features of social-commerce influencers enhance users' trust in the online community and "
  "strengthen their online purchasing intentions.","https://www.frontiersin.org/articles/10.3389/fpsyg.2022.853168"),
 ("A study on the role of social media influencers in India (2025)","identified information "
  "provision, influence, experience sharing, and convenience as key roles, showing a strong "
  "correlation between these roles and the products consumers actually buy.",
  "https://link.springer.com/10.1007/978-981-96-4319-6_25"),
 ("A study on online influencers in the Indian clothing sector (2026)","found that younger "
  "consumers and students are the most influenced group, that lower-income consumers focus on "
  "discounts and affordability, and that wealthier customers seek brand exclusivity and quality.",
  "https://link.springer.com/10.1007/978-3-032-10016-0_59"),
 ("A study on social media advertising and consumer behaviour (2025)","found that the credibility, "
  "perceived authenticity, and sustainability of social postings enhance consumer purchase "
  "intentions, with trust mediating the relationship.",
  "https://www.frontiersin.org/journals/communication/articles/10.3389/fcomm.2025.1595796/full"),
 ("A study titled 'From scroll to sale' (2025)","found that social media triggers significantly "
  "enhance social media influence, which predicts consumer buying behaviour, while influencer "
  "credibility showed no significant influence, indicating possible influencer fatigue.",
  "https://www.frontiersin.org/journals/communication/articles/10.3389/fcomm.2025.1664694/full"),
 ("A study on user-generated content and purchase behaviour (2025)","using consumption-value "
  "theory, found that perceived value from user-generated content and attitude toward it positively "
  "affect consumer purchase behaviour.","https://www.tandfonline.com/doi/full/10.1080/23311975.2025.2471528"),
 ("An empirical study on social media advertising effectiveness among Indian millennials","found "
  "that informativeness, entertainment, credibility, and incentives significantly predict attitudes "
  "toward social media advertising, which in turn predict purchase intention.",
  "https://www.igi-global.com/article/an-empirical-study/249190"),
 ("A systematic review of influencer marketing in the fashion industry (2025)","highlighted "
  "credibility, trust, and attractiveness as the most studied antecedents of influencer-marketing "
  "effectiveness.","https://www.frontiersin.org/journals/communication/articles/10.3389/fcomm.2025.1676901/full"),
 ("An empirical study on influencer attributes in social media (2025)","found that visual "
  "aesthetics and inspiration significantly influence social satisfaction and, in turn, purchase "
  "intentions.","https://www.nature.com/articles/s41598-025-03336-6"),
 ("A study on influencer marketing and e-commerce (2026)","observed that consumers show cautious "
  "and often neutral trust toward influencer promotions, but content consistency, transparency, and "
  "perceived expertise are associated with stronger purchase intentions.",
  "https://www.frontiersin.org/journals/communication/articles/10.3389/fcomm.2026.1723356/full"),
]
for i,(a,f,u) in enumerate(lit,1):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    run=p.add_run(f"{i}. {a} "); run.bold=True; p.add_run(f+" ")
    s=p.add_run(f"(Source: {u})"); s.italic=True; s.font.size=Pt(9)
heading("2.1 Research Gap",2)
para("The literature confirms that social media, influencers, advertisements, and user-generated "
     "content strongly influence consumer purchase intention. However, many studies focus on a "
     "single dimension (such as influencers or advertising) or on purchase intention rather than "
     "self-reported buying behaviour, and relatively few test whether the overall influence differs "
     "across gender and income groups within a single Indian sample. The present study addresses "
     "this gap by measuring the combined influence of multiple social media factors on buying "
     "behaviour among 286 respondents and testing for demographic differences using a t-test and "
     "ANOVA.")
pb()

# ================= CHAPTER 3
heading("CHAPTER 3: RESEARCH METHODOLOGY")
heading("3.1 Research Design",2)
para("The study adopts a descriptive research design, aiming to describe and analyse the influence "
     "of social media on consumer buying behaviour and to test specific hypotheses about "
     "demographic differences.")
heading("3.2 Sources of Data",2)
para("Primary data were collected through a structured questionnaire administered via Google Forms. "
     "Secondary data were obtained from journals, research articles, and reliable online sources as "
     "presented in the review of literature.")
heading("3.3 Sampling Design and Ethics",2)
para("A convenience sampling method was used. A total of 321 responses were received. The "
     "questionnaire began with an informed-consent question, and only the 286 respondents who "
     "voluntarily agreed to participate were included in the analysis; the 35 who did not consent "
     "were excluded in keeping with research ethics. The effective sample size is therefore 286.")
heading("3.4 Tools for Data Collection",2)
para("The questionnaire captured demographic details (age, gender, education, occupation, income), "
     "social media usage (platform, daily time, purchase history, product category), and a "
     "21-statement attitudinal scale measuring the influence of social media on buying behaviour, "
     "each rated on a five-point Likert scale from 1 (Strongly Disagree) to 5 (Strongly Agree).")
heading("3.5 Tools for Analysis",2)
for t in ["Percentage analysis - to describe the respondent profile.",
    "Mean and standard deviation - to summarise the attitudinal statements.",
    "Independent-samples t-test - to compare the influence score between male and female respondents.",
    "One-way ANOVA - to compare the influence score across income groups (and, supplementary, age groups).",
    "Correlation - to examine the relationship between daily social media usage time and influence.",
    "Cronbach's alpha - to assess the reliability of the scale."]:
    bullet(t)
para("Analysis was carried out using Python (pandas and scipy), and charts were produced using "
    "matplotlib.")
heading("3.6 Hypotheses of the Study",2)
para("Hypothesis 1 (T-test):",bold=True)
para("H0: There is no significant difference in the influence of social media on buying behaviour "
     "between male and female respondents.")
para("H1: There is a significant difference between male and female respondents.")
para("Hypothesis 2 (ANOVA):",bold=True)
para("H0: There is no significant difference in the influence of social media on buying behaviour "
     "across income groups.")
para("H1: There is a significant difference across income groups.")
pb()

# ================= CHAPTER 4
heading("CHAPTER 4: DATA ANALYSIS AND INTERPRETATION")
para("This chapter presents the analysis of the 286 valid responses, organised into percentage "
     "analysis, descriptive analysis of the attitudinal statements, reliability analysis, and "
     "hypothesis testing.")

heading("4.1 Percentage Analysis of Respondent Profile",2)
para("Table 4.1 Classification by Age",bold=True)
table(["Age Group","Respondents","%"],[["18-24",125,43.7],["25-34",70,24.5],["45-54",27,9.4],
    ["Below 18",25,8.7],["35-44",24,8.4],["55 and above",15,5.2],["Total",286,100.0]])
img("s_age.png",4.4,"Figure 4.1 Respondents by Age")
para("Interpretation: The majority of respondents (43.7%) belong to the 18-24 age group, followed "
     "by 25-34 (24.5%), indicating a predominantly young respondent base.")

para("Table 4.2 Classification by Gender",bold=True)
table(["Gender","Respondents","%"],[["Male",150,52.4],["Female",106,37.1],
    ["Prefer not to say",30,10.5],["Total",286,100.0]])
img("s_gender.png",4.0,"Figure 4.2 Respondents by Gender")
para("Interpretation: 52.4% of respondents are male, 37.1% female, and 10.5% preferred not to "
     "disclose their gender.")

para("Table 4.3 Classification by Educational Qualification",bold=True)
table(["Qualification","Respondents","%"],[["Undergraduate",117,40.9],["Postgraduate",75,26.2],
    ["Diploma",37,12.9],["High School",29,10.1],["Others",18,6.3],["Doctorate",10,3.5],["Total",286,100.0]])

para("Table 4.4 Classification by Occupation",bold=True)
table(["Occupation","Respondents","%"],[["Student",119,41.6],["Private Employee",80,28.0],
    ["Self-employed",26,9.1],["Homemaker",19,6.6],["Government Employee",17,5.9],
    ["Others",15,5.2],["Business Owner",10,3.5],["Total",286,100.0]])

para("Table 4.5 Classification by Monthly Income",bold=True)
table(["Monthly Income","Respondents","%"],[["Below ₹20,000",111,38.8],["₹20,001-₹40,000",79,27.6],
    ["₹40,001-₹60,000",47,16.4],["Above ₹80,000",26,9.1],["₹60,001-₹80,000",23,8.0],["Total",286,100.0]])
img("s_income.png",4.4,"Figure 4.3 Respondents by Monthly Income")
para("Interpretation: A majority (38.8%) earn below ₹20,000 per month, consistent with the large "
     "proportion of students in the sample.")

para("Table 4.6 Most Used Social Media Platform",bold=True)
table(["Platform","Respondents","%"],[["Instagram",121,42.3],["YouTube",48,16.8],["Facebook",35,12.2],
    ["WhatsApp",28,9.8],["Snapchat",17,5.9],["LinkedIn",15,5.2],["Others",12,4.2],["X (Twitter)",10,3.5],["Total",286,100.0]])
img("s_platform.png",4.6,"Figure 4.4 Most Used Platform")
para("Interpretation: Instagram is the most used platform (42.3%), followed by YouTube (16.8%), "
     "confirming the visual, content-driven nature of social commerce.")

para("Table 4.7 Daily Time Spent on Social Media",bold=True)
table(["Daily Time","Respondents","%"],[["2-4 hours",101,35.3],["1-2 hours",69,24.1],
    ["Less than 1 hour",43,15.0],["4-6 hours",40,14.0],["More than 6 hours",33,11.5],["Total",286,100.0]])

para("Table 4.8 Have You Purchased After Seeing a Product on Social Media?",bold=True)
table(["Response","Respondents","%"],[["Yes",221,77.3],["No",65,22.7],["Total",286,100.0]])
img("s_purchased.png",4.0,"Figure 4.5 Purchase After Seeing on Social Media")
para("Interpretation: A substantial 77.3% of respondents have purchased a product after seeing it "
     "on social media, demonstrating the strong commercial impact of these platforms.")

para("Table 4.9 Most Purchased Product Category",bold=True)
table(["Product Category","Respondents","%"],[["Fashion & Clothing",85,29.7],["Electronics",67,23.4],
    ["Beauty & Personal Care",44,15.4],["Food & Beverages",39,13.6],["Home & Lifestyle",21,7.3],
    ["Others",16,5.6],["Health & Fitness",14,4.9],["Total",286,100.0]])
para("Interpretation: Fashion & Clothing (29.7%) and Electronics (23.4%) are the most purchased "
     "categories through social media.")

heading("4.2 Descriptive Analysis of Influence Statements",2)
para("Table 4.10 Mean Scores of Attitudinal Statements",bold=True)
items = [
 ["Social media influences my purchasing decisions",3.73,1.14],
 ["Advertisements on social media encourage me to buy products",3.75,1.20],
 ["Influencer recommendations increase my interest in products",3.59,1.23],
 ["I trust products recommended by social media influencers",3.59,1.29],
 ["Customer reviews on social media influence my buying decisions",3.60,1.15],
 ["Product images and videos on social media attract me",3.62,1.19],
 ["Social media promotions and discounts motivate me to purchase",3.65,1.22],
 ["Limited-time offers encourage me to buy immediately",3.63,1.23],
 ["Celebrity endorsements influence my buying decisions",3.76,1.25],
 ["Likes, comments, and shares increase my trust in a product",3.77,1.15],
 ["I compare products after seeing them on social media",3.81,1.21],
 ["Social media helps me make informed purchasing decisions",3.72,1.15],
 ["I follow brands on social media for product updates",3.67,1.19],
 ["Social media has changed my shopping habits",3.66,1.24],
 ["I spend more time exploring products because of social media",3.69,1.20],
 ["I prefer buying products recommended by influencers",3.71,1.17],
 ["Social media provides reliable product information",3.68,1.21],
 ["I feel confident purchasing products promoted on social media",3.60,1.23],
 ["Social media helps me compare prices before buying",3.65,1.28],
 ["User-generated content influences my purchase decisions",3.71,1.18],
 ["Overall, social media has a significant impact on my buying behaviour",3.62,1.17],
]
table(["Statement","Mean","S.D."], items)
para("Interpretation: All statements record mean scores above the neutral value of 3, ranging from "
     "3.59 to 3.81. The highest agreement is for 'I compare products after seeing them on social "
     "media' (3.81) and 'Likes, comments, and shares increase my trust in a product' (3.77). The "
     "overall mean influence score is 3.68, indicating that social media has an above-average "
     "influence on consumer buying behaviour.")

heading("4.3 Reliability Analysis",2)
para("Table 4.11 Cronbach's Alpha",bold=True)
table(["Scale","No. of Items","Cronbach's Alpha"],[["Social media influence",21,0.533]])
para("Interpretation: The Cronbach's alpha for the 21-item scale is 0.533, indicating a moderate "
     "level of internal consistency. While below the ideal threshold of 0.70, this value is "
     "acceptable for exploratory research; the inferential results are nonetheless interpreted with "
     "appropriate caution, and this is acknowledged as a limitation.")

heading("4.4 Hypothesis Testing - Independent Samples T-Test",2)
para("Objective: To test whether the influence of social media on buying behaviour differs "
     "significantly between male and female respondents. (Respondents who preferred not to disclose "
     "their gender were excluded from this two-group comparison.)")
para("Table 4.12 T-Test - Influence Score by Gender",bold=True)
table(["Gender","N","Mean","S.D."],[["Male",150,3.663,0.391],["Female",106,3.759,0.306]])
table(["t-value","p-value","Significance (alpha = 0.05)"],[["-2.201","0.029","Significant"]])
para("Interpretation: The calculated t-value is -2.201 with a p-value of 0.029, which is less than "
     "0.05. Therefore, the null hypothesis is REJECTED. There is a statistically significant "
     "difference in the influence of social media on buying behaviour between male and female "
     "respondents, with female respondents (mean = 3.76) being more influenced than male "
     "respondents (mean = 3.66).")

heading("4.5 Hypothesis Testing - One-Way ANOVA",2)
para("Objective: To test whether the influence of social media differs significantly across income "
     "groups.")
para("Table 4.13 ANOVA - Influence Score by Income",bold=True)
table(["Income Group","N","Mean","S.D."],[["Below ₹20,000",111,3.719,0.327],
    ["₹20,001-₹40,000",79,3.757,0.369],["₹40,001-₹60,000",47,3.570,0.379],
    ["₹60,001-₹80,000",23,3.619,0.423],["Above ₹80,000",26,3.505,0.447]])
table(["F-value","p-value","Significance (alpha = 0.05)"],[["3.854","0.005","Significant"]])
para("Interpretation: The calculated F-value is 3.854 with a p-value of 0.005, which is less than "
     "0.05. Therefore, the null hypothesis is REJECTED. There is a statistically significant "
     "difference in the influence of social media across income groups. The influence is highest "
     "among lower-income respondents (below ₹40,000) and lowest among the highest-income group "
     "(above ₹80,000), suggesting that lower-income consumers are more responsive to social media, "
     "promotions, and discounts.")
img("s_score_by_income.png",4.6,"Figure 4.6 Mean Influence Score by Income Group")

heading("4.6 Supplementary ANOVA - Influence by Age",2)
para("A supplementary one-way ANOVA was conducted across age groups. The result (F = 0.720, "
     "p = 0.609) is greater than 0.05, so the null hypothesis is accepted: there is no significant "
     "difference in the influence of social media on buying behaviour across age groups. The "
     "influence is fairly uniform across all ages in the sample.")
img("s_score_by_age.png",4.6,"Figure 4.7 Mean Influence Score by Age Group")

heading("4.7 Correlation Analysis",2)
para("A correlation was computed between daily social media usage time and the influence score. "
     "The Spearman correlation coefficient is 0.046 (p = 0.436), indicating a negligible and "
     "statistically non-significant relationship. In other words, simply spending more hours on "
     "social media does not, by itself, lead to a stronger influence on buying behaviour; the "
     "quality and nature of engagement appear to matter more than the quantity of time spent.")
pb()

# ================= CHAPTER 5
heading("CHAPTER 5: FINDINGS, SUGGESTIONS, CONCLUSION AND SUMMARY")
heading("5.1 Major Findings",2)
for f in [
 "Of 321 responses received, 286 (89%) provided informed consent and were analysed.",
 "The sample is predominantly young (43.7% aged 18-24) and student-heavy (41.6%).",
 "Instagram is the most used platform (42.3%), followed by YouTube (16.8%).",
 "A large majority (77.3%) have purchased a product after seeing it on social media.",
 "Fashion & Clothing (29.7%) and Electronics (23.4%) are the most purchased categories.",
 "The overall mean influence score is 3.68, indicating an above-average influence of social media on buying behaviour.",
 "All 21 attitudinal statements scored above the neutral value of 3.",
 "The t-test showed a significant difference in influence between genders (t = -2.201, p = 0.029), with females more influenced than males.",
 "The ANOVA showed a significant difference across income groups (F = 3.854, p = 0.005), with lower-income consumers more influenced.",
 "There was no significant difference in influence across age groups (F = 0.720, p = 0.609).",
 "Daily time spent on social media had a negligible correlation with influence (rho = 0.046).",
 "The scale reliability was moderate (Cronbach's alpha = 0.53), noted as a limitation."]:
    bullet(f)
heading("5.2 Suggestions",2)
for s in [
 "Marketers should design visually rich, review-driven campaigns, as product comparison and social proof (likes, shares, reviews) scored highest.",
 "Given the stronger influence on female consumers, brands targeting women can prioritise influencer and social media channels.",
 "Since lower-income consumers are more influenced and respond to discounts, affordability-focused promotions are effective for this segment.",
 "Premium brands targeting higher-income consumers should emphasise quality and exclusivity rather than relying solely on social media promotions.",
 "Consumers should critically evaluate influencer recommendations and verify information before making purchases.",
 "Platforms can promote transparency in sponsored content to maintain consumer trust."]:
    bullet(s)
heading("5.3 Conclusion",2)
para("The study set out to examine the impact of social media on consumer buying behaviour. The "
     "findings clearly demonstrate that social media exerts an above-average influence on buying "
     "decisions, with more than three-quarters of respondents having purchased a product after "
     "encountering it on social media. Importantly, this influence is not uniform across all "
     "consumers: it is significantly stronger among female and lower-income respondents, while it "
     "does not differ significantly across age groups. Product comparison, social proof, and "
     "advertisements emerged as the most influential factors, whereas the sheer amount of time "
     "spent on social media was not a meaningful driver. Although the moderate scale reliability "
     "warrants some caution, the overall evidence confirms that social media is a powerful and "
     "differentiated force in shaping consumer behaviour, and businesses should tailor their "
     "strategies to specific demographic segments accordingly.")
heading("5.4 Summary",2)
para("This project studied the impact of social media on consumer buying behaviour using primary "
     "data from 286 consenting respondents out of 321 surveyed. The data were analysed through "
     "percentage analysis, descriptive statistics, an independent-samples t-test, one-way ANOVA, "
     "and correlation. The results showed a strong overall influence of social media on buying "
     "behaviour, a significant gender difference (females more influenced), a significant income "
     "difference (lower-income consumers more influenced), and no significant age difference. Based "
     "on the findings, suggestions were offered for marketers and consumers.")
pb()

# ================= APPENDIX I
heading("APPENDIX I: QUESTIONNAIRE")
para("Consent: Do you voluntarily agree to participate in this research? (Yes / No)",bold=True)
para("Section A: Demographic and Usage Information",bold=True)
for q in ["1. Age: Below 18 / 18-24 / 25-34 / 35-44 / 45-54 / 55 and above",
 "2. Gender: Male / Female / Prefer not to say",
 "3. Highest Educational Qualification: High School / Diploma / Undergraduate / Postgraduate / Doctorate / Others",
 "4. Occupation: Student / Private Employee / Government Employee / Self-employed / Business Owner / Homemaker / Others",
 "5. Monthly Income: Below ₹20,000 / ₹20,001-₹40,000 / ₹40,001-₹60,000 / ₹60,001-₹80,000 / Above ₹80,000",
 "6. Which social media platform do you use most?",
 "7. How much time do you spend on social media daily?",
 "8. Have you ever purchased a product after seeing it on social media? (Yes / No)",
 "9. Which type of products do you purchase most through social media?"]:
    para(q)
para("Section B: Influence of Social Media on Buying Behaviour (1 = Strongly Disagree to 5 = Strongly Agree)",bold=True)
for i,(txt,_,_) in enumerate(items,10):
    para(f"{i}. {txt}")
pb()

# ================= APPENDIX II
heading("APPENDIX II: BIBLIOGRAPHY / REFERENCES")
para("Note: Final formatting should be adjusted to the citation style (APA/MLA) required by your "
     "institution.", italic=True)
for i,(a,_,u) in enumerate(lit,1):
    p=doc.add_paragraph(); p.add_run(f"{i}. {a}. Retrieved from {u}").font.size=Pt(10)

out=os.path.join(OUT_DIR,"Social_Media_Research_Report.docx")
doc.save(out)
print("Saved:", out)
