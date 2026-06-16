#!/usr/bin/env python3
"""
Generates a polished, plain-English Word guide for the StreamerClips branding
repository: what the assets are, the colour palette, how the logo generator
(render_logo.py) works, and how to use / re-render the marks.

Reuses the document style/approach of the original generate_guide.py.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Colour palette (mirrors the StreamerClips brand)
# ----------------------------------------------------------------------------
NAVY = RGBColor(0x0C, 0x12, 0x1F)   # near-black background top
BLUE = RGBColor(0x1A, 0x68, 0xFF)   # electric blue (deep)
CYAN = RGBColor(0x60, 0xD2, 0xFF)   # electric blue (light)
TEAL = RGBColor(0x40, 0xC4, 0xFF)   # frame ring
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT = RGBColor(0x7A, 0x7A, 0x7A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(text, level=1, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(18 if level == 1 else 14 if level == 2 else 12)
    return p


def add_para(text, bold=False, italic=False, color=None, size=11, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_numbered(text, bold_lead=None):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_callout(label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "EAF6FF")
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    r = cell.paragraphs[0].add_run(label)
    r.bold = True
    r.font.color.rgb = BLUE
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(text)
    r2.italic = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_code(lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "0C121F")
    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = CYAN
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def make_table(headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], "0C121F")
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        shade = "FFFFFF" if r_i % 2 == 0 else "EAF6FF"
        for i, val in enumerate(row):
            set_cell_bg(cells[i], shade)
            para = cells[i].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9.5)
            if i == 0:
                run.bold = True
    return table


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C9D4E2")
    pBdr.append(bottom)
    pPr.append(pBdr)


def page_break():
    doc.add_page_break()


# ============================================================================
# COVER
# ============================================================================
for _ in range(4):
    doc.add_paragraph()

add_para("BRAND ASSET HANDBOOK", bold=True, color=TEAL, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("StreamerClips")
r.bold = True
r.font.size = Pt(34)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Branding & Logo Guide")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BLUE

add_para("An SC monogram with an integrated play button \u2014 premium black + electric blue.",
         color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

add_para(
    "A plain-English guide to the StreamerClips brand assets: what every file is, "
    "the colour palette, how the logo is generated in pure Python, and how to "
    "re-render or tweak it yourself.",
    italic=True, color=LIGHT, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ============================================================================
# OVERVIEW
# ============================================================================
add_heading("1. What This Repository Is", level=1)
add_para(
    "This repository (klingesh/Klingesh-streamerclips) holds the brand assets for "
    "StreamerClips, a gaming / streaming brand. The centrepiece is a logo: a "
    "minimalist SC monogram where the letter C frames a play triangle \u2014 the "
    "universal 'watch / clip' symbol \u2014 inside a circular badge. It is designed "
    "to look great as a YouTube profile picture when cropped to a circle.")
add_para(
    "Everything is delivered in two forms: scalable vector files (SVG) that can be "
    "exported at any size without losing quality, and ready-to-use raster images "
    "(PNG) at 800 \u00d7 800 pixels.")

add_callout(
    "In plain English:",
    "It is a logo kit. You get the StreamerClips badge in a glowing 'neon' style "
    "and a clean 'flat' style, each as both an editable vector and a ready-to-post "
    "image \u2014 plus the Python script that builds them all.")

add_heading("Two styles to choose from", level=2, color=BLUE)
add_bullet(" a neon edge-glow for a futuristic, eye-catching look.", bold_lead="Glow \u2013")
add_bullet(" cleaner and more minimal, no glow \u2014 better for small sizes and simple placements.", bold_lead="Flat \u2013")

add_heading("Repository layout", level=2, color=BLUE)
add_code([
    "Klingesh-streamerclips/",
    "  README.md                  <- top-level overview",
    "  .gitignore",
    "  branding/",
    "    README.md                <- asset details + palette",
    "    render_logo.py           <- the logo generator (pure Python)",
    "    streamerclips_logo.svg / .png            (glow, dark badge)",
    "    streamerclips_logo_transparent.png       (glow, transparent)",
    "    streamerclips_logo_flat.svg / .png       (flat, dark badge)",
    "    streamerclips_logo_flat_transparent.png  (flat, transparent)",
])

page_break()

# ============================================================================
# FILES
# ============================================================================
add_heading("2. The Files \u2014 What Each One Is For", level=1)
add_para("All raster (PNG) files are 800 \u00d7 800 pixels. The SVG files are "
         "resolution-independent, so you can export any size you need from them.",
         space_after=8)

make_table(
    ["File", "Background", "Style", "Best used for"],
    [
        ["streamerclips_logo.png", "Dark badge", "Glow", "YouTube avatar (neon look)"],
        ["streamerclips_logo_transparent.png", "Transparent", "Glow", "Overlays / watermarks"],
        ["streamerclips_logo.svg", "Dark badge", "Glow", "Vector master (glow)"],
        ["streamerclips_logo_flat.png", "Dark badge", "Flat", "YouTube avatar (minimal)"],
        ["streamerclips_logo_flat_transparent.png", "Transparent", "Flat", "Overlays / watermarks"],
        ["streamerclips_logo_flat.svg", "Dark badge", "Flat", "Vector master (flat)"],
    ],
)

add_para("", space_after=4)
add_callout(
    "Which one should I use?",
    "Posting a profile picture? Use a .png with a dark badge. Placing the logo on "
    "top of a video, thumbnail, or coloured background? Use a _transparent.png. "
    "Need a poster, banner, or any large/print size? Start from a .svg and export "
    "at the size you want.")

add_heading("Vector (SVG) vs. raster (PNG) \u2014 quick reminder", level=2, color=BLUE)
add_bullet(" math-based shapes that stay razor-sharp at any size. Edit in tools like Illustrator/Inkscape, or export to any resolution. This is your 'master' copy.", bold_lead="SVG \u2013")
add_bullet(" a fixed grid of pixels (here 800\u00d7800). Ready to upload immediately, but it gets blurry if blown up much larger.", bold_lead="PNG \u2013")

page_break()

# ============================================================================
# PALETTE
# ============================================================================
add_heading("3. The Colour Palette", level=1)
add_para("The brand uses a premium 'black + electric blue' scheme. Here are the "
         "exact colours used in the logo:", space_after=8)

make_table(
    ["Role", "Hex code", "What it's used for"],
    [
        ["Background top", "#0C121F", "Top of the dark badge gradient"],
        ["Background bottom", "#03050B", "Bottom of the dark badge gradient"],
        ["Electric blue (light)", "#60D2FF", "Top of the SC letter gradient"],
        ["Electric blue (deep)", "#1A68FF", "Bottom of the SC letter gradient"],
        ["Play triangle (light)", "#F0FAFF", "Top of the play-button gradient"],
        ["Frame ring", "#40C4FF", "The thin circular ring around the badge"],
    ],
)

add_para("", space_after=4)
add_callout(
    "Why gradients?",
    "Instead of one flat colour, the letters and background fade smoothly from one "
    "shade to another (top to bottom). This gives the mark depth and that polished, "
    "'lit-from-within' premium feel.")

page_break()

# ============================================================================
# HOW THE GENERATOR WORKS
# ============================================================================
add_heading("4. How the Logo Is Made (render_logo.py)", level=1)
add_para(
    "The logo is not drawn by hand in a design app \u2014 it is generated by code. "
    "render_logo.py is a single, pure-Python script (no third-party libraries "
    "required) that draws the mark mathematically and writes out all six files.")

add_callout(
    "The clever bit:",
    "The script describes the logo as math (circles, arcs, and a triangle) rather "
    "than pixels. Because the shapes are formulas, the same definitions produce both "
    "a crisp vector SVG and a smooth, anti-aliased PNG \u2014 and changing a number "
    "instantly changes the logo everywhere.")

add_heading("The building blocks of the mark", level=2, color=BLUE)
add_bullet(" two round-capped curved strokes (arcs) that join in the middle to form the letter shape.", bold_lead="The 'S' \u2013")
add_bullet(" a single curved stroke (an arc left open on the right side) that forms the letter and 'cups' the play button.", bold_lead="The 'C' \u2013")
add_bullet(" a rounded triangle pointing right \u2014 the classic 'play / watch' symbol \u2014 nestled inside the C.", bold_lead="The play triangle \u2013")
add_bullet(" a thin circle around everything, tying the badge together.", bold_lead="The frame ring \u2013")

add_heading("How it actually draws (in simple terms)", level=2, color=BLUE)
add_para("The script uses a technique called a 'signed distance field.' That sounds "
         "technical, but the idea is simple:")
add_numbered(" For every pixel in the 800\u00d7800 image, it asks: 'how far am I from each shape?'", bold_lead="Measure distance \u2013")
add_numbered(" If a pixel is inside a shape it gets that shape's colour; right at the edge it gets a soft blend (this is anti-aliasing, which removes jagged edges).", bold_lead="Decide colour \u2013")
add_numbered(" Letters fade from light to deep blue, the triangle fades to near-white, and the background fades from dark to darker.", bold_lead="Apply gradients \u2013")
add_numbered(" In glow mode, pixels near the mark get extra blue light added, creating the neon halo. Flat mode simply skips this step.", bold_lead="Add the glow \u2013")
add_numbered(" The finished pixels are packed into PNG files; the same shapes are also written out as SVG paths.", bold_lead="Save files \u2013")

add_heading("What you can safely tweak", level=2, color=BLUE)
add_para("All the knobs live near the top of the script, clearly labelled:")
add_bullet(" change any hex value to recolour the mark.", bold_lead="Palette (colours) \u2013")
add_bullet(" centre points, radii, and stroke widths control the size, spacing, and thickness of the S, C, ring, and triangle.", bold_lead="Geometry \u2013")
add_bullet(" passing glow=True or glow=False switches the neon halo on or off.", bold_lead="Glow toggle \u2013")

page_break()

# ============================================================================
# HOW TO RE-RENDER
# ============================================================================
add_heading("5. How to Re-render or Change the Logo", level=1)
add_para("Because the logo is just a Python script with no dependencies, re-creating "
         "all six files takes one command.")

add_heading("Step 1 \u2014 Run the generator", level=2, color=BLUE)
add_code([
    "cd branding",
    "python3 render_logo.py",
])
add_para("This overwrites the two PNGs and one SVG for the glow style, plus the "
         "matching three files for the flat style \u2014 all six in one go.", space_after=8)

add_heading("Step 2 \u2014 Change something (optional)", level=2, color=BLUE)
add_para("Open render_logo.py and edit the values near the top. A few examples:")
add_bullet(" change BLUE_HI / BLUE_LO (the letter gradient) to a different colour pair.", bold_lead="Recolour the letters \u2013")
add_bullet(" increase the S_H and C_H half-stroke-width values.", bold_lead="Make strokes thicker \u2013")
add_bullet(" adjust RING_R (radius) and RING_HW (thickness).", bold_lead="Resize the ring \u2013")
add_para("Then run the script again to see your changes in all the output files.", space_after=8)

add_callout(
    "Good practice:",
    "Make one small change at a time and re-render, so you can clearly see what each "
    "value does. Keep the SVG files as your master \u2014 they stay sharp at any size.")

add_heading("Step 3 \u2014 Use the files", level=2, color=BLUE)
add_bullet(" upload streamerclips_logo.png (or the flat version) as your channel avatar.", bold_lead="YouTube / social avatar \u2013")
add_bullet(" use a _transparent.png so the badge sits cleanly on any background.", bold_lead="Overlays & watermarks \u2013")
add_bullet(" open or export from a .svg for banners, posters, or merchandise.", bold_lead="Large / print use \u2013")

page_break()

# ============================================================================
# GLOSSARY
# ============================================================================
add_heading("6. Plain-English Glossary", level=1)
glossary = [
    ("Monogram", "A design made from letters \u2014 here the 'S' and 'C' of StreamerClips."),
    ("Raster / PNG", "An image made of a fixed grid of pixels; great for posting, blurry if enlarged too much."),
    ("Vector / SVG", "A shape described by math; stays perfectly sharp at any size and is easy to edit."),
    ("Transparent background", "No solid backdrop, so the logo sits cleanly on top of anything."),
    ("Gradient", "A smooth fade between two colours, used here for depth and a premium look."),
    ("Anti-aliasing", "Softening the edges of shapes so they look smooth instead of jagged."),
    ("Glow / neon edge", "Soft blue light added around the mark for a futuristic effect."),
    ("Signed distance field", "The maths trick the script uses to know how far each pixel is from a shape."),
    ("Hex code", "A six-character colour code like #1A68FF that names an exact colour."),
    ("Dependency-free", "The script needs only standard Python \u2014 nothing extra to install."),
]
for term, meaning in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(term + ": ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(meaning)

divider()
add_para(
    "In short: this repo is a self-contained logo kit. Use the PNGs to post right "
    "away, the SVGs for anything big, and render_logo.py whenever you want to "
    "recolour, resize, or restyle the StreamerClips mark.",
    italic=True, color=GREY, size=11)

doc.save("/projects/sandbox/Klingesh-streamerclips/branding/StreamerClips_Branding_Guide.docx")
print("Branding guide created successfully.")
