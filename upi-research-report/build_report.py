# -*- coding: utf-8 -*-
"""
Builds the full research report (Word .docx) from the REAL analysis results.
Topic: Influence of Digital Payment Apps (UPI) on Personal Spending and Saving Behaviour
"""
import os, shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/projects/sandbox"
OUT_DIR = "/projects/sandbox/Klingesh-streamerclips/upi-research-report"
os.makedirs(OUT_DIR, exist_ok=True)
for png in ["r_gender.png","r_age.png","r_income.png","r_app.png","r_saving_by_income.png"]:
    shutil.copy(os.path.join(SRC, png), os.path.join(OUT_DIR, png))

doc = Document()
# Base style
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)

def add_page_break():
    doc.add_page_break()

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    return h

def para(text, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    doc.add_paragraph(text, style="List Number")

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.size = Pt(11)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(11)
    doc.add_paragraph()
    return t

def add_image(path, width=5.2, caption=None):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(10)

# ============================================================ TITLE PAGE
ti = doc.add_paragraph(); ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ti.add_run("INFLUENCE OF DIGITAL PAYMENT APPS (UPI) ON\nPERSONAL SPENDING AND SAVING BEHAVIOUR")
r.bold = True; r.font.size = Pt(18)
doc.add_paragraph()
for line, sz, bold in [
    ("A PROJECT REPORT", 13, True),
    ("Submitted in partial fulfilment of the requirements", 12, False),
    ("for the award of the degree of", 12, False),
    ("MASTER OF BUSINESS ADMINISTRATION (FINANCE)", 13, True),
    ("", 12, False),
    ("Submitted by", 12, False),
    ("[Your Name]", 13, True),
    ("[Register Number]", 12, False),
    ("", 12, False),
    ("Under the guidance of", 12, False),
    ("[Guide Name]", 13, True),
    ("", 12, False),
    ("[Department / School Name]", 12, True),
    ("[University / College Name]", 12, True),
    ("June 2026", 12, True),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(line); rr.bold = bold; rr.font.size = Pt(sz)
add_page_break()

# ============================================================ DECLARATION
heading("DECLARATION", 1)
para('I hereby declare that the project report titled "Influence of Digital Payment Apps (UPI) '
     'on Personal Spending and Saving Behaviour" submitted by me is a record of original work '
     'carried out under the guidance of [Guide Name]. The findings reported in this study are '
     'based on primary data collected through a structured questionnaire and have not been '
     'submitted earlier for the award of any degree, diploma, or similar title.')
doc.add_paragraph(); doc.add_paragraph()
para("Place: [City]"); para("Date: [Date]")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run("[Your Name]\n[Register Number]").bold = True
add_page_break()

# ============================================================ ACKNOWLEDGEMENT
heading("ACKNOWLEDGEMENT", 1)
para("I express my sincere gratitude to my guide [Guide Name] for the valuable guidance, "
     "encouragement, and support extended throughout this study. I am thankful to the Head of "
     "the Department and the faculty members of [Department Name] for their constant motivation. "
     "I also thank the 171 respondents who spared their valuable time to complete the "
     "questionnaire, without whom this study would not have been possible. Finally, I am grateful "
     "to my family and friends for their continuous support and encouragement.")
add_page_break()

# ============================================================ ABSTRACT
heading("ABSTRACT", 1)
para("The Unified Payments Interface (UPI) has transformed the way individuals in India transact, "
     "save, and manage money. This study examines the influence of UPI-based digital payment "
     "applications on the personal spending and saving behaviour of users. Primary data were "
     "collected from 171 respondents through a structured questionnaire measured on a five-point "
     "Likert scale. The data were analysed using percentage analysis, an independent-samples "
     "t-test, one-way ANOVA, and Pearson correlation. The results indicate that respondents "
     "report a high level of UPI usage for daily transactions, with a moderate perceived influence "
     "on spending (mean = 3.32) and a relatively higher positive perception of saving discipline "
     "(mean = 3.80). The t-test found no statistically significant difference in spending behaviour "
     "between male and female users, and ANOVA found no statistically significant difference in "
     "saving behaviour across income groups. The study concludes that while UPI increases "
     "transactional convenience and encourages routine digital record-keeping, its perceived "
     "influence on spending and saving is broadly similar across demographic groups. "
     "Recommendations for promoting responsible digital payment habits are offered.")
para("Keywords: UPI, digital payments, spending behaviour, saving behaviour, financial discipline, "
     "fintech.", italic=True)
add_page_break()

# ============================================================ TABLE OF CONTENTS
heading("TABLE OF CONTENTS", 1)
toc = [
    ("Chapter 1", "Introduction"),
    ("Chapter 2", "Review of Literature"),
    ("Chapter 3", "Research Methodology"),
    ("Chapter 4", "Data Analysis and Interpretation"),
    ("Chapter 5", "Findings, Suggestions, Conclusion and Summary"),
    ("", "Appendix I - Questionnaire"),
    ("", "Appendix II - Bibliography / References"),
]
make_table(["Chapter", "Title"], toc)
add_page_break()

# ============================================================ CHAPTER 1
heading("CHAPTER 1: INTRODUCTION", 1)

heading("1.1 Background of the Study", 2)
para("The financial landscape of India has undergone a remarkable transformation over the past "
     "decade. The introduction of the Unified Payments Interface (UPI) by the National Payments "
     "Corporation of India (NPCI) in 2016 created a real-time, mobile-first payment system that "
     "allows users to transfer money instantly between bank accounts using a smartphone. What "
     "began as a convenient alternative to cash and cards has rapidly become the dominant mode of "
     "retail payment in the country. By 2025, UPI was processing a very large share of the world's "
     "real-time digital transactions, reflecting an unprecedented level of adoption across age "
     "groups, income segments, and geographies.")
para("Applications such as Google Pay, PhonePe, Paytm and BHIM have made digital payments "
     "effortless. A transaction that once required carrying cash, visiting an ATM, or swiping a "
     "card can now be completed in seconds by scanning a QR code. This frictionless experience has "
     "important behavioural consequences. When money becomes intangible and payment becomes "
     "instantaneous, the psychological 'pain of paying' is reduced, which can influence how much "
     "and how often people spend, as well as how disciplined they are about saving.")

heading("1.2 Concept of Digital Payments and UPI", 2)
para("Digital payment refers to a transaction in which value is transferred electronically without "
     "the physical exchange of currency. UPI is a layer that sits on top of the existing banking "
     "infrastructure and enables peer-to-peer and person-to-merchant transactions through a "
     "single mobile application linked to a Virtual Payment Address (VPA). The key features of "
     "UPI include 24x7 availability, instant settlement, low or zero transaction cost for users, "
     "interoperability across banks and apps, and a simple authentication process.")

heading("1.3 Spending and Saving Behaviour", 2)
para("Spending behaviour refers to the patterns, frequency, and decision-making processes that "
     "govern how individuals use their money for consumption. Saving behaviour refers to the "
     "extent to which individuals set aside a portion of their income for future needs and "
     "exercise financial discipline. Both behaviours are influenced by income, financial literacy, "
     "convenience of payment, social influence, and the availability of incentives such as "
     "cashbacks and discounts. The convenience offered by UPI can pull behaviour in two opposite "
     "directions: it may encourage impulsive and more frequent spending, while at the same time "
     "the automatic digital record of every transaction may help users track expenses and budget "
     "more effectively.")

heading("1.4 Statement of the Problem", 2)
para("While UPI has clearly increased the convenience and speed of payments, its behavioural "
     "impact on personal finance is not fully understood. There is a concern that the ease of "
     "digital payments may lead to increased and unplanned spending, weaker awareness of how much "
     "is being spent, and a possible decline in saving. At the same time, digital records and "
     "budgeting features may improve financial discipline. It is therefore important to study "
     "empirically how UPI influences the spending and saving behaviour of users and whether this "
     "influence differs across gender and income groups.")

heading("1.5 Objectives of the Study", 2)
for o in [
    "To study the usage pattern of UPI-based digital payment applications among respondents.",
    "To examine the influence of UPI on the personal spending behaviour of users.",
    "To analyse the influence of UPI on the personal saving behaviour of users.",
    "To study whether spending behaviour differs significantly between male and female users.",
    "To study whether saving behaviour differs significantly across income groups.",
    "To offer suggestions for the responsible use of digital payment applications.",
]:
    numbered(o)

heading("1.6 Scope of the Study", 2)
para("The study focuses on individual users of UPI-based digital payment applications. It covers "
     "the behavioural dimensions of spending and saving as perceived and self-reported by the "
     "respondents. The study is based on primary data collected from 171 respondents through an "
     "online structured questionnaire. The geographic scope is limited to the respondents reached "
     "through the survey, and the study period covers the data collection conducted in 2026.")

heading("1.7 Need and Significance of the Study", 2)
para("As UPI becomes the default mode of payment for a large section of the population, "
     "understanding its behavioural consequences is valuable for several stakeholders. For "
     "individuals, the study highlights how digital payments may affect their own financial "
     "discipline. For policymakers and financial educators, it provides insight into the need for "
     "financial-literacy initiatives in a cashless environment. For fintech companies, it offers a "
     "perspective on designing features that promote responsible spending and saving.")

heading("1.8 Limitations of the Study", 2)
for l in [
    "The study relies on self-reported perceptions, which may differ from actual transaction data.",
    "The sample is dominated by younger respondents (below 25 years) and lower income groups, which may limit generalisability.",
    "The reliability of the attitudinal scales, as discussed in Chapter 4, was found to be low, and the findings should be interpreted with this in mind.",
    "The study is cross-sectional and does not capture changes in behaviour over time.",
]:
    bullet(l)
add_page_break()

# ============================================================ CHAPTER 2
heading("CHAPTER 2: REVIEW OF LITERATURE", 1)
para("This chapter reviews relevant studies published over the past ten years on digital payments, "
     "UPI, and their influence on consumer spending and saving behaviour. The review helps to "
     "establish the theoretical foundation of the study and to identify the research gap.")

lit = [
    ("Agarwal, Ghosh, Li and Ruan (2018)",
     "studied the impact of the 2016 demonetization in India and found that the usage of digital "
     "payments rose and monthly spending increased among households that were previously more "
     "cash-dependent, with spending remaining elevated even after cash availability recovered.",
     "https://papers.ssrn.com/sol3/papers.cfm?abstract_id=3641508"),
    ("Brown, Nacht, Nellen and Stix (2023)",
     "examined cashless payments and consumer spending and found that present-biased consumers "
     "tend to spend more the more frequently they use cashless payment instruments, suggesting a "
     "behavioural link between payment convenience and higher expenditure.",
     "https://papers.ssrn.com/sol3/papers.cfm?abstract_id=4668928"),
    ("A study on UPI's impact on spending behaviour among Indian users (2024)",
     "surveyed 276 respondents and reported that approximately 75% experienced increased spending "
     "due to UPI, attributing this largely to the intangible nature of digital money which reduced "
     "the guilt usually associated with spending.",
     "https://arxiv.org/abs/2401.09937"),
    ("A study on e-wallets and impulse buying among Gen-Y and Gen-Z consumers (2022)",
     "found that perceived enjoyment of using an e-wallet positively influenced impulse buying, "
     "while satisfaction with the e-wallet did not have a significant relationship with impulse "
     "purchases.",
     "https://link.springer.com/article/10.1057/s41270-022-00164-9"),
    ("A study on sales promotion, FOMO and digital payment methods (2024)",
     "highlighted that despite their convenience, digital payment methods can contribute to "
     "impulse purchases, particularly among the younger generation influenced by fear of missing "
     "out and promotional offers.",
     "https://www.tandfonline.com/doi/abs/10.1080/23311975.2024.2419484"),
    ("A study on the impact of instant digital transactions on consumer spending (2025)",
     "observed that the rise of instant digital transactions through UPI, mobile wallets and "
     "contactless payments enhanced convenience but also encouraged impulsive purchases and "
     "contributed to reduced savings.",
     "https://www.ijraset.com/research-paper/impact-of-instant-digital-transactions-on-consumer-spending-behaviour"),
    ("A study on the behavioural impact of digital wallets on modern consumers (2024)",
     "analysed how digital wallets have changed consumer expenditure trends and concluded that "
     "convenience and security significantly shape purchasing behaviour and consumption patterns.",
     "https://www.researchgate.net/publication/383360479"),
    ("A study on determinants of behavioural intention to use digital payment among Indian "
     "youngsters (2024)",
     "found that perceived value, trust, compatibility and social influence significantly "
     "influence the behavioural intention to use mobile wallets, whereas perceived enjoyment had a "
     "weaker effect.",
     "https://www.mdpi.com/1911-8074/17/2/87"),
    ("A study on UPI as a digital innovation and financial inclusion (2021)",
     "found that UPI positively impacts financial literacy, which in turn significantly influences "
     "financial inclusion and contributes to economic development.",
     "https://www.researchgate.net/publication/353016519"),
    ("A study on cashless transactions among university students (2021)",
     "reported that cashless payment modes are widely adopted by students and that incentives such "
     "as perceived free benefits increase their willingness to spend more.",
     "https://pmc.ncbi.nlm.nih.gov/articles/PMC10719559/"),
    ("A comparative analysis of UPI versus traditional payment methods (2024)",
     "examined efficiency and user experience and concluded that UPI's seamless, real-time system "
     "offers clear advantages over cash and cards and promotes a cashless economy.",
     "https://www.researchgate.net/publication/384286994"),
    ("An empirical study on consumer perception towards UPI in India (2024)",
     "explored consumer perception of UPI as an instant inter-bank payment system and identified "
     "convenience and trust as key drivers of adoption.",
     "https://link.springer.com/chapter/10.1007/978-3-031-64359-0_13"),
    ("A study on intention and adoption of e-wallets (2022)",
     "found that perceived usefulness, ease of use, social influence, lifestyle compatibility and "
     "trust significantly and positively affect the intention to use and adopt e-wallets.",
     "https://www.researchgate.net/publication/348518228"),
    ("A study on sustainable economic development, digital payment and consumer demand (2022)",
     "established that the growth of digital payments is associated with changes in consumer demand "
     "and contributes to sustainable economic development.",
     "https://www.mdpi.com/1660-4601/19/14/8819"),
    ("An EY-CII industry report on digital payments in rural and semi-urban India (2024)",
     "reported that UPI is the most preferred payment mode for a large proportion of users in "
     "rural and semi-urban areas, and that a very high share of respondents demonstrated a strong "
     "inclination to save and invest.",
     "https://www.ey.com/en_in/newsroom/2024/12"),
    ("A study assessing the role of UPI in financial inclusion and digital trust (2025)",
     "found that financial literacy and infrastructure significantly affect service experience and "
     "digital trust, although trust and policy awareness alone did not significantly drive "
     "sustained usage.",
     "https://www.abacademies.org/articles/assessing-the-role-of-upi-17741.html"),
    ("A study on the influence of digital wallets on consumer behaviour for sustainable growth "
     "(2025)",
     "examined how shifts in consumer purchasing behaviour driven by digital wallets affect "
     "sustainable economic growth.",
     "https://link.springer.com/chapter/10.1007/978-981-95-0599-9_8"),
    ("A study on small-scale merchants' acceptance of UPI in India (2026)",
     "identified a gap in understanding UPI acceptance among small merchants and examined the "
     "factors driving their adoption of UPI payment apps.",
     "https://www.nature.com/articles/s41599-026-08005-1"),
]
for i, (author, finding, url) in enumerate(lit, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(f"{i}. {author} ")
    run.bold = True
    p.add_run(finding + " ")
    src = p.add_run(f"(Source: {url})")
    src.italic = True; src.font.size = Pt(9)

heading("2.1 Research Gap", 2)
para("The reviewed literature confirms that digital payments and UPI strongly influence consumer "
     "behaviour, with many studies reporting increased and impulsive spending. However, most "
     "studies focus either on the adoption of UPI or on spending alone, and relatively few examine "
     "spending and saving behaviour together while also testing whether these effects differ "
     "across gender and income groups. The present study addresses this gap by empirically "
     "analysing both spending and saving behaviour among 171 UPI users and testing for "
     "demographic differences using a t-test and ANOVA.")
add_page_break()

# ============================================================ CHAPTER 3
heading("CHAPTER 3: RESEARCH METHODOLOGY", 1)

heading("3.1 Research Design", 2)
para("The study adopts a descriptive research design, as it aims to describe and analyse the "
     "spending and saving behaviour of UPI users and to test specific hypotheses regarding "
     "demographic differences.")

heading("3.2 Sources of Data", 2)
para("Primary data were collected directly from respondents using a structured questionnaire "
     "administered through Google Forms. Secondary data were drawn from journals, research "
     "articles, industry reports, and reliable online sources, as presented in the review of "
     "literature.")

heading("3.3 Sampling Design", 2)
para("A convenience sampling method was used to reach respondents who actively use UPI "
     "applications. The sample size for the study is 171 respondents.")

heading("3.4 Tools for Data Collection", 2)
para("A structured questionnaire consisting of 25 questions was used. The first part captured "
     "demographic information (gender, age, education, income, and most-used UPI app). The second "
     "part contained ten statements measuring spending behaviour, and the third part contained ten "
     "statements measuring saving behaviour, each rated on a five-point Likert scale ranging from "
     "1 (Strongly Disagree) to 5 (Strongly Agree).")

heading("3.5 Tools for Analysis", 2)
for t in [
    "Percentage analysis - to describe the demographic profile of respondents.",
    "Mean and standard deviation - to summarise responses to the spending and saving statements.",
    "Independent-samples t-test - to compare spending behaviour between two gender groups.",
    "One-way ANOVA - to compare saving behaviour across more than two income groups.",
    "Pearson correlation - to examine the relationship between spending and saving scores.",
    "Cronbach's alpha - to assess the internal reliability of the scales.",
]:
    bullet(t)
para("The analysis was carried out using Python (pandas and scipy libraries), and charts were "
     "generated using the matplotlib library.")

heading("3.6 Hypotheses of the Study", 2)
para("Hypothesis 1 (T-test):", bold=True)
para("H0: There is no significant difference in spending behaviour between male and female UPI users.")
para("H1: There is a significant difference in spending behaviour between male and female UPI users.")
para("Hypothesis 2 (ANOVA):", bold=True)
para("H0: There is no significant difference in saving behaviour across different income groups.")
para("H1: There is a significant difference in saving behaviour across different income groups.")
add_page_break()

# ============================================================ CHAPTER 4
heading("CHAPTER 4: DATA ANALYSIS AND INTERPRETATION", 1)
para("This chapter presents the analysis of the primary data collected from 171 respondents. "
     "It is organised into percentage analysis of the demographic profile, descriptive analysis of "
     "the spending and saving statements, reliability analysis, and hypothesis testing.")

heading("4.1 Percentage Analysis of Demographic Profile", 2)

para("Table 4.1 Classification by Gender", bold=True)
make_table(["Gender", "No. of Respondents", "Percentage (%)"],
           [["Male", 98, 57.3], ["Female", 73, 42.7], ["Total", 171, 100.0]])
add_image(os.path.join(OUT_DIR, "r_gender.png"), 4.2, "Figure 4.1 Respondents by Gender")
para("Interpretation: Out of 171 respondents, 57.3% are male and 42.7% are female, indicating a "
     "fairly balanced participation with a slight majority of male respondents.")

para("Table 4.2 Classification by Age", bold=True)
make_table(["Age Group", "No. of Respondents", "Percentage (%)"],
           [["Below 25", 139, 81.3], ["25-35", 29, 17.0], ["Above 35", 3, 1.8], ["Total", 171, 100.0]])
add_image(os.path.join(OUT_DIR, "r_age.png"), 4.2, "Figure 4.2 Respondents by Age")
para("Interpretation: The sample is dominated by respondents below 25 years (81.3%), reflecting "
     "the strong adoption of UPI among the youth.")

para("Table 4.3 Classification by Educational Qualification", bold=True)
make_table(["Qualification", "No. of Respondents", "Percentage (%)"],
           [["UG", 75, 43.9], ["Professional / Masters", 74, 43.3], ["School", 22, 12.9], ["Total", 171, 100.0]])
para("Interpretation: Respondents are almost evenly split between undergraduate (43.9%) and "
     "professional/masters (43.3%) qualifications, indicating an educated respondent base.")

para("Table 4.4 Classification by Monthly Income", bold=True)
make_table(["Monthly Income (Rs.)", "No. of Respondents", "Percentage (%)"],
           [["Below 25,000", 112, 65.5], ["25,000-50,000", 40, 23.4],
            ["50,000-1,00,000", 16, 9.4], ["Above 1,00,000", 3, 1.8], ["Total", 171, 100.0]])
add_image(os.path.join(OUT_DIR, "r_income.png"), 4.2, "Figure 4.3 Respondents by Monthly Income")
para("Interpretation: A majority of respondents (65.5%) fall in the below Rs.25,000 income group, "
     "consistent with the young, largely student profile of the sample.")

para("Table 4.5 Most Used UPI Application", bold=True)
make_table(["UPI App", "No. of Respondents", "Percentage (%)"],
           [["Google Pay", 77, 45.0], ["PhonePe", 67, 39.2], ["Paytm", 22, 12.9],
            ["Others", 4, 2.3], ["BHIM", 1, 0.6], ["Total", 171, 100.0]])
add_image(os.path.join(OUT_DIR, "r_app.png"), 4.2, "Figure 4.4 Most Used UPI App")
para("Interpretation: Google Pay (45.0%) and PhonePe (39.2%) together account for more than "
     "four-fifths of the respondents, confirming their market dominance.")

heading("4.2 Descriptive Analysis of Spending Behaviour Statements", 2)
para("Table 4.6 Mean Scores - Spending Behaviour", bold=True)
spend_items = [
    ["I use UPI for most of my daily transactions", 4.30, 0.98],
    ["Using UPI makes me spend more than I planned", 3.12, 1.21],
    ["I make more impulse/unplanned purchases because UPI is convenient", 3.27, 1.16],
    ["I find it harder to track my spending when I use UPI", 2.52, 1.14],
    ["UPI cashbacks and offers encourage me to spend more", 3.53, 1.03],
    ["I spend on small things now that I would have skipped with cash", 3.51, 1.15],
    ["The ease of UPI has increased my overall monthly expenses", 3.17, 1.05],
    ["I often pay digitally without thinking about the amount", 2.98, 1.15],
    ["I shop online more because UPI payment is quick and easy", 3.63, 1.15],
    ["I lose the feel of money when I pay through UPI compared to cash", 3.17, 1.25],
]
make_table(["Statement", "Mean", "S.D."], spend_items)
para("Interpretation: The highest mean is for 'I use UPI for most of my daily transactions' "
     "(4.30), confirming heavy UPI usage. Statements relating to increased spending due to "
     "cashbacks (3.53), online shopping (3.63), and spending on small items (3.51) record "
     "above-neutral means, indicating a moderate tendency of UPI to encourage spending. The "
     "overall mean spending score is 3.32, suggesting a moderate perceived influence on spending.")

heading("4.3 Descriptive Analysis of Saving Behaviour Statements", 2)
para("Table 4.7 Mean Scores - Saving Behaviour", bold=True)
save_items = [
    ["I am able to save a fixed amount every month", 3.81, 1.08],
    ["Using UPI helps me keep digital records that improve my budgeting", 4.17, 0.95],
    ["I review my UPI transaction history to control my expenses", 3.70, 1.11],
    ["UPI has made me more disciplined about my finances", 3.72, 1.09],
    ["I set a monthly spending limit for myself", 3.55, 1.15],
    ["I transfer money to savings before spending on wants", 2.98, 1.23],
    ["Digital payments help me avoid borrowing/cash shortages", 4.12, 1.06],
    ["I plan my expenses better because of digital transaction records", 3.80, 1.08],
    ["My savings have remained stable or improved since using UPI", 3.84, 1.03],
    ["Overall, UPI has had a positive impact on my saving habits", 4.30, 0.87],
]
make_table(["Statement", "Mean", "S.D."], save_items)
para("Interpretation: Saving-related statements record generally high means. The strongest "
     "agreement is with 'Overall, UPI has had a positive impact on my saving habits' (4.30) and "
     "'Using UPI helps me keep digital records that improve my budgeting' (4.17). The overall "
     "mean saving score is 3.80, which is higher than the spending score, suggesting that "
     "respondents perceive UPI as supporting financial record-keeping and discipline.")

heading("4.4 Reliability Analysis", 2)
para("Table 4.8 Cronbach's Alpha", bold=True)
make_table(["Scale", "No. of Items", "Cronbach's Alpha"],
           [["Spending behaviour", 10, 0.118], ["Saving behaviour", 10, -0.054]])
para("Interpretation: The Cronbach's alpha values for both scales are well below the commonly "
     "accepted threshold of 0.70. This indicates low internal consistency among the scale items, "
     "meaning respondents did not answer the related statements in a highly consistent pattern. "
     "This is acknowledged as a limitation of the study, and the subsequent inferential results "
     "are therefore interpreted with appropriate caution.")

heading("4.5 Hypothesis Testing - Independent Samples T-Test", 2)
para("Objective: To test whether spending behaviour differs significantly between male and female "
     "UPI users.")
para("Table 4.9 T-Test - Spending Score by Gender", bold=True)
make_table(["Gender", "N", "Mean", "S.D."],
           [["Male", 98, 3.297, 0.361], ["Female", 73, 3.352, 0.400]])
make_table(["t-value", "p-value", "Significance (alpha = 0.05)"],
           [["-0.929", "0.354", "Not significant"]])
para("Interpretation: The calculated t-value is -0.929 with a p-value of 0.354, which is greater "
     "than 0.05. Therefore, the null hypothesis is ACCEPTED. There is no statistically significant "
     "difference in spending behaviour between male and female UPI users. Both groups report a "
     "similar, moderate level of UPI-influenced spending.")

heading("4.6 Hypothesis Testing - One-Way ANOVA", 2)
para("Objective: To test whether saving behaviour differs significantly across income groups.")
para("Table 4.10 ANOVA - Saving Score by Income", bold=True)
make_table(["Income Group (Rs.)", "N", "Mean", "S.D."],
           [["Below 25,000", 112, 3.790, 0.333], ["25,000-50,000", 40, 3.810, 0.299],
            ["50,000-1,00,000", 16, 3.737, 0.391], ["Above 1,00,000", 3, 4.167, 0.115]])
make_table(["F-value", "p-value", "Significance (alpha = 0.05)"],
           [["1.472", "0.224", "Not significant"]])
para("Interpretation: The calculated F-value is 1.472 with a p-value of 0.224, which is greater "
     "than 0.05. Therefore, the null hypothesis is ACCEPTED. There is no statistically significant "
     "difference in saving behaviour across the different income groups. Saving behaviour related "
     "to UPI is broadly similar regardless of the respondent's income level.")
add_image(os.path.join(OUT_DIR, "r_saving_by_income.png"), 4.6, "Figure 4.5 Mean Saving Score by Income Group")

heading("4.7 Correlation Analysis", 2)
para("A Pearson correlation was computed between the spending score and the saving score. The "
     "correlation coefficient (r) is 0.059 with a p-value of 0.444, indicating a negligible and "
     "statistically non-significant relationship between perceived spending influence and perceived "
     "saving behaviour. In other words, respondents who report higher UPI-influenced spending do "
     "not necessarily report lower saving, and vice versa.")
add_page_break()

# ============================================================ CHAPTER 5
heading("CHAPTER 5: FINDINGS, SUGGESTIONS, CONCLUSION AND SUMMARY", 1)

heading("5.1 Major Findings", 2)
for f in [
    "Out of 171 respondents, 57.3% are male and 42.7% are female.",
    "The sample is dominated by young users, with 81.3% below 25 years of age.",
    "A majority of respondents (65.5%) earn below Rs.25,000 per month.",
    "Google Pay (45.0%) and PhonePe (39.2%) are the most widely used UPI applications.",
    "UPI is heavily used for daily transactions, with the highest mean score of 4.30.",
    "The overall mean spending score is 3.32, indicating a moderate perceived influence of UPI on spending.",
    "The overall mean saving score is 3.80, indicating a relatively positive perception of UPI's role in saving and budgeting.",
    "The t-test showed no significant difference in spending behaviour between male and female users (t = -0.929, p = 0.354).",
    "The ANOVA showed no significant difference in saving behaviour across income groups (F = 1.472, p = 0.224).",
    "The correlation between spending and saving scores was negligible (r = 0.059).",
    "The reliability (Cronbach's alpha) of both scales was low, which is noted as a limitation.",
]:
    bullet(f)

heading("5.2 Suggestions", 2)
for s in [
    "UPI applications should include built-in monthly spending limits and alerts to help users avoid unplanned spending.",
    "Users should regularly review their UPI transaction history to track and control expenses.",
    "Financial literacy programmes should be introduced for young users to encourage disciplined saving in a cashless environment.",
    "App providers can offer 'savings nudges' such as round-up savings and goal-based reminders alongside cashbacks.",
    "Educational institutions can incorporate basic personal-finance modules that address digital spending habits.",
]:
    bullet(s)

heading("5.3 Conclusion", 2)
para("The study set out to examine the influence of UPI-based digital payment applications on the "
     "personal spending and saving behaviour of users. The findings indicate that UPI is deeply "
     "integrated into the daily transactions of respondents and exerts a moderate influence on "
     "spending, particularly through cashbacks, online shopping, and small everyday purchases. At "
     "the same time, respondents perceive UPI positively in terms of record-keeping and budgeting, "
     "reflected in a higher saving score. Statistically, neither spending behaviour (by gender) "
     "nor saving behaviour (by income) differed significantly, suggesting that the behavioural "
     "influence of UPI is broadly uniform across these demographic groups. While the low "
     "reliability of the scales calls for cautious interpretation, the overall evidence suggests "
     "that UPI is a double-edged tool: it enhances convenience and encourages spending, yet also "
     "supports financial awareness through digital records. Promoting responsible usage through "
     "financial literacy and app-level features is therefore essential.")

heading("5.4 Summary", 2)
para("This project studied the influence of UPI on personal spending and saving behaviour using "
     "primary data from 171 respondents. The data were analysed through percentage analysis, "
     "descriptive statistics, an independent-samples t-test, one-way ANOVA, and correlation. The "
     "results showed heavy UPI usage, a moderate influence on spending, a positive perception of "
     "saving and budgeting, and no significant demographic differences in either behaviour. Based "
     "on the findings, suggestions were offered to encourage responsible digital payment habits.")
add_page_break()

# ============================================================ APPENDIX I
heading("APPENDIX I: QUESTIONNAIRE", 1)
para("Section A: Demographic Information", bold=True)
for q in [
    "1. Gender: Male / Female",
    "2. Age: Below 25 / 25-35 / Above 35",
    "3. Educational Qualification: School / UG / PG / Professional",
    "4. Monthly Income: Below Rs.25,000 / Rs.25,000-50,000 / Rs.50,000-1,00,000 / Above Rs.1,00,000",
    "5. Which UPI app do you use the most? Google Pay / PhonePe / Paytm / BHIM / Others",
]:
    para(q)
para("Section B: Spending Behaviour (1 = Strongly Disagree to 5 = Strongly Agree)", bold=True)
for i, (txt, _, _) in enumerate(spend_items, 6):
    para(f"{i}. {txt}")
para("Section C: Saving Behaviour (1 = Strongly Disagree to 5 = Strongly Agree)", bold=True)
for i, (txt, _, _) in enumerate(save_items, 16):
    para(f"{i}. {txt}")
add_page_break()

# ============================================================ APPENDIX II
heading("APPENDIX II: BIBLIOGRAPHY / REFERENCES", 1)
para("Note: The following sources were referred to during the study. Final formatting should be "
     "adjusted to the citation style (APA/MLA) required by your institution.", italic=True)
for i, (author, _, url) in enumerate(lit, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. {author}. Retrieved from {url}").font.size = Pt(10)

out_path = os.path.join(OUT_DIR, "UPI_Research_Report.docx")
doc.save(out_path)
print("Saved:", out_path)
EOF_MARKER = None
